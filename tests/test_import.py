import shutil
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from transcript_toolkit.errors import ToolkitError
from transcript_toolkit.core.tables import untimed_ids
from transcript_toolkit.project import init_project
from transcript_toolkit.steps.import_ import (run_import, run_import_unsynced,
                                              timestamp_regimes)

FIXTURES = Path(__file__).parent / "fixtures"


def make_docx(project, name: str, lines: list[str]) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(project.data_dir / name)


PER_PARAGRAPH = [                                   # every paragraph carries its own [HH:MM:SS]
    "[00:00:05] Q: Where did you grow up?",
    "[00:00:12] Delta: In a small town by the sea.",
    "[00:01:04] My father was a fisherman and my mother taught school.",
    "[00:02:20] Q: And what came after school?",
    "[00:02:30] Delta: University, then a long detour into journalism.",
]
PER_TURN_ONLY = [                                   # timestamp only on each speaker turn
    "[00:00:05] Q: Where did you grow up?",
    "[00:00:12] Echo: In a small town by the sea.",
    "My father was a fisherman and my mother taught school.",
    "That shaped how I saw the world for a long time afterwards.",
    "[00:02:20] Q: And what came after school?",
    "[00:02:30] Echo: University, then a long detour into journalism.",
]
FIXTURE_DOCX = [
    "Fake_Alpha_20240101_session1_SYNC.docx",
    "Fake_Alpha_20240108_session2_SYNC.docx",
    "Fake, Beta_SYNC.docx",
]


@pytest.fixture
def project(tmp_path):
    project = init_project(str(tmp_path / "ws"))
    for name in FIXTURE_DOCX:
        shutil.copy(FIXTURES / name, project.data_dir / name)
    return project


def test_import_end_to_end(project, capsys):
    df = run_import(project)
    assert project.paragraphs_path.exists()
    assert project.paragraphs_path.with_suffix(".csv").exists()
    assert set(df["interview_id"]) == {
        "fake_alpha_20240101_session1", "fake_alpha_20240108_session2", "fake_beta"}
    on_disk = pd.read_parquet(project.paragraphs_path)
    assert len(on_disk) == len(df) == 13 + 7 + 9 - 1   # docx paragraphs minus the orphan
    log = (project.logs_dir / "import_warnings.log").read_text()
    assert "before the first speaker turn" in log      # the orphan section
    assert "kept with the current speaker" in log      # the benign colon-note section
    out = capsys.readouterr().out
    assert "3 transcripts" in out
    assert "fake_alpha" in out and "session1" in out   # narrator-pooling table printed


def test_all_per_paragraph_no_timestamp_warning(project, capsys):
    make_docx(project, "Fake_Delta_SYNC.docx", PER_PARAGRAPH)
    # remove the mixed fixtures so only the clean transcript is present
    for f in project.data_dir.glob("Fake_Alpha*"):
        f.unlink()
    (project.data_dir / "Fake, Beta_SYNC.docx").unlink()
    df = run_import(project)
    assert all(r["ok"] for r in timestamp_regimes(df))
    out = capsys.readouterr().out
    assert "every paragraph carries its own" in out
    assert "⚠ Timestamps" not in out


def test_per_turn_only_transcript_warns(project, capsys):
    make_docx(project, "Fake_Echo_SYNC.docx", PER_TURN_ONLY)
    for f in project.data_dir.glob("Fake_Alpha*"):
        f.unlink()
    (project.data_dir / "Fake, Beta_SYNC.docx").unlink()
    df = run_import(project)
    echo = next(r for r in timestamp_regimes(df) if r["interview_id"] == "fake_echo")
    assert echo["coverage"] == 0.0 and echo["n_cont"] == 2 and not echo["ok"]
    out = capsys.readouterr().out
    assert "⚠ Timestamps" in out
    assert "speaker turns only" in out
    assert "fake_echo" in out


def test_import_is_idempotent(project):
    df1 = run_import(project)
    df2 = run_import(project)
    pd.testing.assert_frame_equal(df1, df2)


def test_duplicate_ids_abort(project):
    # same interview id from a second filename ("Fake, Beta" vs "Fake_Beta" + other suffix)
    shutil.copy(FIXTURES / "Fake, Beta_SYNC.docx", project.data_dir / "Fake_Beta_final.docx")
    with pytest.raises(ToolkitError, match="same interview id"):
        run_import(project)


def test_timestampless_docx_aborts_with_hint(project):
    doc = Document()
    doc.add_paragraph("Q: This transcript has no timestamps at all.")
    doc.add_paragraph("Gamma: So the parser must reject it loudly.")
    doc.save(project.data_dir / "Fake_Gamma_SYNC.docx")
    with pytest.raises(ToolkitError, match=r"Fake_Gamma_SYNC\.docx"):
        run_import(project)


def test_no_docx_aborts(tmp_path):
    project = init_project(str(tmp_path / "empty-ws"))
    with pytest.raises(ToolkitError, match="No .docx transcripts"):
        run_import(project)


def test_word_lock_files_ignored(project):
    shutil.copy(FIXTURES / "Fake, Beta_SYNC.docx", project.data_dir / "~$ke, Beta_SYNC.docx")
    df = run_import(project)
    assert set(df["interview_id"]) == {
        "fake_alpha_20240101_session1", "fake_alpha_20240108_session2", "fake_beta"}


# --- transcripts that were never SYNC'd -------------------------------------------------------

UNSYNCED = FIXTURES / "unsynced" / "Fake_Gamma_Transcript.docx"


def put_unsynced(project, name: str = "Fake_Gamma_Transcript.docx"):
    project.unsynced_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy(UNSYNCED, project.unsynced_dir / name)


def test_untimed_transcripts_join_the_one_collection(project, capsys):
    """They used to be parsed into a table of their own, because a summary was all that could
    be made from them. A clip is a run of paragraphs, and paragraph numbers are something every
    transcript has — so they are imported with the rest and go through every step."""
    put_unsynced(project)
    df = run_import(project)

    assert "fake_gamma_transcript" in set(df["interview_id"])
    gamma = df[df["interview_id"] == "fake_gamma_transcript"]
    assert (gamma["turn_time_start"] == "").all() and (gamma["sub_time_start"] == "").all()
    assert set(gamma["speaker_role"]) == {"Interviewer", "Narrator"}
    # the sentence with a colon in it continues the turn rather than starting one
    assert (gamma["turn_idx"].value_counts() > 1).any()
    assert untimed_ids(df) == {"fake_gamma_transcript"}
    assert "never SYNC'd" in capsys.readouterr().out


def test_an_untimed_transcript_is_not_reported_as_a_broken_one(project, capsys):
    """It has no timestamps by definition. Measuring its timestamp coverage would report the
    thing that is true of it as a fault in it, on every single import."""
    put_unsynced(project)
    run_import(project)
    out = capsys.readouterr().out
    assert "fake_gamma_transcript" not in out.split("⚠ Timestamps:")[-1].split("\n\n")[0]
    from transcript_toolkit.steps.import_ import timestamp_regimes
    import pandas as pd
    regimes = timestamp_regimes(pd.read_parquet(project.paragraphs_path))
    assert "fake_gamma_transcript" not in {r["interview_id"] for r in regimes}


def test_the_front_matter_is_left_out_and_written_down(project):
    """A title page and a preface are about the interview, not part of it. Dropping them
    silently would leave nobody able to check what was dropped."""
    put_unsynced(project)
    df = run_import(project)
    assert not df["speech"].str.contains("PREFACE").any()
    log = (project.logs_dir / "import_warnings.log").read_text()
    assert "PREFACE" in log and "Front matter" in log


def test_a_transcript_with_no_times_in_the_wrong_folder_still_fails_loudly(project):
    """`data/` is where a SYNC'd transcript goes, and one with no times in it is a mistake —
    a corrupted export, a wrong file. Reading it as untimed would hide that. The message says
    which folder it belongs in if it really has no times."""
    shutil.copy(UNSYNCED, project.data_dir / "Fake_Gamma_Transcript.docx")
    with pytest.raises(ToolkitError, match="unsynced"):
        run_import(project)


def test_the_same_narrator_cannot_be_in_both_folders(project):
    """Sessions are pooled by narrator, so the same person arriving from both piles would be
    two half-interviews claiming one row."""
    put_unsynced(project, "Fake, Beta.docx")           # the same narrator key as the SYNC'd one
    with pytest.raises(ToolkitError, match="both transcript folders"):
        run_import(project)


def test_the_unsynced_flag_still_works_and_says_it_is_the_same_import(project, capsys):
    """Somebody following an older note types it. It must not look like the untimed folder was
    skipped this time."""
    put_unsynced(project)
    df = run_import_unsynced(project)
    assert "fake_gamma_transcript" in set(df["interview_id"])
    assert "importing both" in capsys.readouterr().out


def test_the_old_separate_table_is_cleared_away(project):
    """A project imported by an earlier toolkit has one. Two answers to "what is in this
    project?" is one too many, and the docx files are the one that matters."""
    from transcript_toolkit.state import load_state, save_state

    put_unsynced(project)
    run_import(project)
    project.unsynced_paragraphs_path.write_bytes(b"old")        # as an older import left it
    state = load_state(project)
    state["steps"]["summarize:unsynced"] = {"full": {"n_units": 1}}
    save_state(project, state)

    run_import(project)
    assert not project.unsynced_paragraphs_path.exists()
    assert "summarize:unsynced" not in load_state(project)["steps"]


def test_both_folders_are_the_collection_in_status(project):
    """One import reads both, so both are counted and either being newer means there is an
    import to run."""
    from transcript_toolkit.steps.status import gather_status

    put_unsynced(project)
    run_import(project)
    status = gather_status(project)
    assert status["docx_files"] == 4 and status["unsynced_files"] == 1
    assert status["imported"] and not status["import_stale"]

    put_unsynced(project, "Another_Untimed.docx")
    assert gather_status(project)["import_stale"]
